import io
import os
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from backend.core.logging import logger
from fastapi import HTTPException, status

class DocumentExtractor:
    """Extracts text content and page metadata from supported document formats."""

    @staticmethod
    def extract_from_pdf(contents: bytes) -> List[Dict[str, Any]]:
        """Extract text page by page from PDF files."""
        pages = []
        try:
            reader = PdfReader(io.BytesIO(contents))
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({
                        "page_number": idx + 1,
                        "text": text
                    })
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse PDF document: {str(e)}"
            )
        if not pages:
            pages.append({"page_number": 1, "text": ""})
        return pages

    @staticmethod
    def extract_from_docx(contents: bytes) -> List[Dict[str, Any]]:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(io.BytesIO(contents))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            full_text = "\n\n".join(paragraphs)
            return [{"page_number": 1, "text": full_text}]
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse DOCX document: {str(e)}"
            )

    @staticmethod
    def extract_from_txt(contents: bytes) -> List[Dict[str, Any]]:
        """Extract text from plain text or Markdown files."""
        try:
            text = contents.decode("utf-8", errors="replace")
            return [{"page_number": 1, "text": text}]
        except Exception as e:
            logger.error(f"Error extracting TXT/MD: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse text document: {str(e)}"
            )

    @staticmethod
    def extract_from_html(contents: bytes) -> List[Dict[str, Any]]:
        """Extract cleaned text from HTML documents."""
        try:
            html_text = contents.decode("utf-8", errors="replace")
            soup = BeautifulSoup(html_text, "html.parser")
            # Remove scripts, styles, metadata
            for element in soup(["script", "style", "meta", "noscript", "header", "footer"]):
                element.decompose()
            text = soup.get_text(separator="\n\n")
            return [{"page_number": 1, "text": text}]
        except Exception as e:
            logger.error(f"Error extracting HTML: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse HTML document: {str(e)}"
            )

    @classmethod
    def extract(cls, filename: str, contents: bytes) -> List[Dict[str, Any]]:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return cls.extract_from_pdf(contents)
        elif ext == ".docx":
            return cls.extract_from_docx(contents)
        elif ext in [".txt", ".md"]:
            return cls.extract_from_txt(contents)
        elif ext in [".html", ".htm"]:
            return cls.extract_from_html(contents)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file format: {ext}"
            )
